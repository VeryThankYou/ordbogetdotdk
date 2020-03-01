import docx, random
from wordAPI import api



def fuckUpDoc(file):
    document = docx.Document(str(file))
    fullText = []
    #Først åbnes det valgte dokument, og der oprettes en liste der kan indeholde teksten fra paragraferne i dokumentet
    for paragraph in document.paragraphs:
        fullText.append(paragraph.text)
    #Her læses paragraferne, og de gemmes i listen
    stringtext = " ".join(fullText)

    listOfWords = stringtext.split()
    #Først dannes en fuld string af al teksten med join(), hvor der sættes mellemrum mellem elementerne i listen. Derefter splittes alle ordene op i en ny liste, hvor der splittes ved hvert mellemrum
    newList = []
    #Her oprettes en liste der skal indeholde alle ord, som indeholder bogstaver, og alle special-tegn og tal slettes fra disse ord
    for e in listOfWords:
        n=0
        #For hvert ord i den første liste sættes et tal, n, til 0
        for letter in e:
            if letter.isalpha() == 1:
                n += 1
                #For hvert tegn, hvis det er et bogstav, lægges 1 til n
        if n>0:
            e = ''.join(l for l in e if l.isalpha())
            newList.append(e)
            n=0
            #Hvis n stadig er 0 til sidst, tages ordet, alle tal og special-tegn slettes, og det tilføjes listen med faktiske ord.


    numsyn = int(len(newList)/50)

    wordchoice = random.choices(newList, k=numsyn)
    #Der vælges nu 2% af ordene tilfældigt

    synonyms = api(wordchoice)
    #Nu oprettes en liste med et synonym for hvert af de tilfældigt valgte ord

    for i in range(len(wordchoice)):
        for e in range(len(listOfWords)):
            try:
                if wordchoice[i-1] in listOfWords[e-1]:
                    listOfWords[e-1] = synonyms[i-1]
                    #For hvert af de tilfældigt valgte ord, tjekkes nu om det indeholdes af et ord i den første liste af ord, inden tegn blev frasorteret. Hvis dette er tilfældet, sættes dette ord fra den originale liste til synonymet til ordet
                    break
            except IndexError:
                continue
            #I tilfælde af indexfejl med mængden af synonymer foregår dette i try- except-sætning
            

    newText = ' '.join(listOfWords)
    #Nu sættes alle ordene sammen med mellemrum

    newdoc = docx.Document()
    newdoc.add_paragraph(newText)
    #Der oprettes nu et nyt dokument, hvor al den nye tekst skrives i en paragraf, uden nogen form for formatering

    document.save(str(file))
    newdoc.save(str(file))
    #Nu overwrites den originale fil med den nye tekst
